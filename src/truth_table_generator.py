import re
from itertools import product
from sympy import symbols, sympify
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor


# ==============================
# Function Definitions
# ==============================

def parse_boolean_expression(expr_input: str):
    """
    Convert user input Boolean expression to Python/sympy compatible format.
    
    Supports:
      - ¬ as NOT
      - * or ^ as AND
      - + or ∨ as OR
      - Variable names can include letters, digits, and underscores (e.g., x1, x_2, F)
    
    Returns:
      sympy_expr: sympy expression object
      variables: list of sympy symbols (variables)
      vars_names: sorted list of variable names as strings

    Raises:
      ValueError: if the input is invalid or contains no variables
    """
    expr_input = expr_input.strip()  # remove leading/trailing spaces

    if not expr_input:
        raise ValueError("Expression cannot be empty.")

    try:
        # Replace user symbols with Python/sympy operators
        expr_str = (expr_input.replace('¬', '~')
                               .replace('*', '&')
                               .replace('^', '&')
                               .replace('+', '|')
                               .replace('∨', '|'))

        # Extract variable names (letters, digits, underscores)
        vars_names = sorted(set(re.findall(r'\b[A-Za-z_][A-Za-z0-9_]*\b', expr_str)))
        if not vars_names:
            raise ValueError("No valid variables found in expression (expected names like x1, x_2, F).")

        # Create sympy symbols
        variables = symbols(vars_names)
        var_dict = {str(v): v for v in variables}

        # Parse into sympy expression
        sympy_expr = sympify(expr_str, locals=var_dict)

        return sympy_expr, variables, vars_names

    except Exception as e:
        raise ValueError(f"Invalid Boolean expression: '{expr_input}'.\nError: {e}")


def generate_truth_table(sympy_expr, variables):
    """
    Generate truth table data for given sympy expression and variable list.
    Returns list of rows: each row = [set_number, var_values..., function_value]
    """
    rows_data = []
    for i, values in enumerate(product([0, 1], repeat=len(variables))):
        val_dict = {variables[j]: values[j] for j in range(len(variables))}
        func_val = int(bool(sympy_expr.subs(val_dict)))
        rows_data.append([str(i)] + [str(v) for v in values] + [str(func_val)])
    return rows_data


def create_docx_table(rows_data, vars_names, user_expression, output_file="truth_table.docx"):
    """
    Create a DOCX table with visible borders from truth table data.
    All cells are horizontally and vertically centered.
    Adds a paragraph at the top with the original Boolean expression.
    Font: Times New Roman, size 12.
    """
    doc = Document()
    
    # Add paragraph with the user's original expression
    para = doc.add_paragraph()
    run = para.add_run(f"Boolean function: {user_expression}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center the paragraph

    table = doc.add_table(rows=len(rows_data) + 1, cols=len(vars_names) + 2)

    # Helper function to center cell text and set font
    def set_cell_center(cell, text):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # horizontal
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')  # vertical
        tcPr.append(vAlign)

    # Add header row
    hdr_cells = table.rows[0].cells
    set_cell_center(hdr_cells[0], "№")
    for j, name in enumerate(vars_names):
        set_cell_center(hdr_cells[j + 1], name)
    set_cell_center(hdr_cells[-1], "f")

    # Add data rows
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            set_cell_center(table.rows[i + 1].cells[j], val)

    # Set visible borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    doc.save(output_file)
    print(f"Truth table saved to {output_file}")


# ==============================
# Main Execution
# ==============================
def main():
    """
    Main program loop: repeatedly ask for Boolean expression until valid,
    then generate truth table and save it as DOCX.
    """
    while True:
        user_input = input("Enter Boolean expression (e.g., ¬(F*¬G)+¬(C+D)+E+F*G*¬H): ").strip()
        if not user_input:
            print("Error: Expression cannot be empty. Please try again.\n")
            continue
        try:
            expr, variables, vars_names = parse_boolean_expression(user_input)
            break  # Valid input, exit loop
        except ValueError as e:
            print(f"Error: {e}\nPlease try again.\n")

    # Generate truth table
    rows_data = generate_truth_table(expr, variables)

    # Create DOCX table
    create_docx_table(rows_data, vars_names, user_input)


if __name__ == "__main__":
    main()
